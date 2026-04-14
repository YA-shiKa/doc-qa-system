"""
utils/document_parser.py — Multi-format document parsing with table extraction.

Handles:
  - PDF (text + scanned via OCR + table extraction via pdfplumber)
  - DOCX
  - Plain text / Markdown

Returns a unified ParsedDocument with structured sections.
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import pdfplumber
import pytesseract
from PIL import Image
from docx import Document as DocxDocument
from core.logging import get_logger

logger = get_logger(__name__)


@dataclass
class DocumentSection:
    """A logical chunk of a document with rich metadata."""
    content: str
    section_title: Optional[str] = None
    page_number: Optional[int] = None
    section_type: str = "text"           # "text" | "table" | "heading" | "caption"
    bbox: Optional[tuple[float, ...]] = None
    confidence: float = 1.0              # OCR confidence if applicable

    @property
    def word_count(self) -> int:
        return len(self.content.split())


@dataclass
class ParsedDocument:
    """Fully parsed document with all sections and metadata."""
    doc_id: str
    filename: str
    title: Optional[str] = None
    author: Optional[str] = None
    sections: list[DocumentSection] = field(default_factory=list)
    total_pages: int = 0
    language: str = "en"
    file_type: str = "unknown"
    metadata: dict = field(default_factory=dict)

    @property
    def full_text(self) -> str:
        return "\n\n".join(s.content for s in self.sections)

    @property
    def tables(self) -> list[DocumentSection]:
        return [s for s in self.sections if s.section_type == "table"]

    @property
    def text_sections(self) -> list[DocumentSection]:
        return [s for s in self.sections if s.section_type in ("text", "heading")]


class DocumentParser:
    """
    Multi-format parser combining direct text extraction with OCR fallback.

    Strategy:
    1. Try native text extraction (fast, accurate)
    2. If page has no extractable text → OCR via pytesseract
    3. Extract tables separately via pdfplumber
    4. Tag each section with page + type metadata
    """

    MIN_TEXT_LENGTH = 20  # Below this, assume scanned page → trigger OCR

    def parse(self, file_path: Path, doc_id: str) -> ParsedDocument:
        """Parse any supported document type."""
        suffix = file_path.suffix.lower()
        logger.info("parsing_document", doc_id=doc_id, file_type=suffix)

        if suffix == ".pdf":
            return self._parse_pdf(file_path, doc_id)
        elif suffix in (".docx", ".doc"):
            return self._parse_docx(file_path, doc_id)
        elif suffix in (".txt", ".md"):
            return self._parse_text(file_path, doc_id)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

    def _parse_pdf(self, path: Path, doc_id: str) -> ParsedDocument:
        sections: list[DocumentSection] = []
        metadata: dict = {}

        with pdfplumber.open(path) as pdf:
            total_pages = len(pdf.pages)
            metadata = pdf.metadata or {}
            title = metadata.get("Title") or metadata.get("title")

            for page_num, page in enumerate(pdf.pages, start=1):
                # 1. Extract native text
                raw_text = page.extract_text() or ""

                if len(raw_text.strip()) < self.MIN_TEXT_LENGTH:
                    # 2. Fallback to OCR
                    raw_text, confidence = self._ocr_page(page)
                    section_type = "text"
                else:
                    confidence = 1.0
                    section_type = "text"

                # Split into logical paragraphs
                paragraphs = self._split_paragraphs(raw_text)
                for para in paragraphs:
                    if not para.strip():
                        continue
                    sec_type = "heading" if self._looks_like_heading(para) else section_type
                    sections.append(DocumentSection(
                        content=para.strip(),
                        page_number=page_num,
                        section_type=sec_type,
                        confidence=confidence,
                    ))

                # 3. Extract tables
                tables = page.extract_tables()
                for table in tables:
                    table_text = self._table_to_markdown(table)
                    if table_text:
                        sections.append(DocumentSection(
                            content=table_text,
                            page_number=page_num,
                            section_type="table",
                        ))

        return ParsedDocument(
            doc_id=doc_id,
            filename=path.name,
            title=title,
            author=metadata.get("Author"),
            sections=sections,
            total_pages=total_pages,
            file_type="pdf",
            metadata=metadata,
        )

    def _parse_docx(self, path: Path, doc_id: str) -> ParsedDocument:
        doc = DocxDocument(path)
        sections: list[DocumentSection] = []

        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            is_heading = para.style.name.startswith("Heading")
            sections.append(DocumentSection(
                content=para.text.strip(),
                section_type="heading" if is_heading else "text",
            ))

        # Extract tables from DOCX
        for table in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            table_text = self._table_to_markdown(rows)
            if table_text:
                sections.append(DocumentSection(content=table_text, section_type="table"))

        # Extract title from core properties
        title = None
        try:
            title = doc.core_properties.title or None
        except Exception:
            pass

        return ParsedDocument(
            doc_id=doc_id,
            filename=path.name,
            title=title,
            sections=sections,
            file_type="docx",
        )

    def _parse_text(self, path: Path, doc_id: str) -> ParsedDocument:
        content = path.read_text(encoding="utf-8", errors="replace")
        paragraphs = self._split_paragraphs(content)
        sections = [
            DocumentSection(content=p.strip(), section_type="text")
            for p in paragraphs if p.strip()
        ]
        return ParsedDocument(
            doc_id=doc_id,
            filename=path.name,
            sections=sections,
            file_type=path.suffix.lstrip("."),
        )

    def _ocr_page(self, page) -> tuple[str, float]:
        """Render PDF page to image and apply OCR."""
        try:
            img = page.to_image(resolution=200).original
            text = pytesseract.image_to_string(img, lang="eng")
            # Estimate confidence from pytesseract data
            data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
            confidences = [int(c) for c in data["conf"] if str(c).isdigit() and int(c) >= 0]
            confidence = sum(confidences) / len(confidences) / 100 if confidences else 0.5
            return text, confidence
        except Exception as e:
            logger.warning("ocr_failed", error=str(e))
            return "", 0.0

    def _split_paragraphs(self, text: str) -> list[str]:
        """Split text into paragraphs by double newline or section breaks."""
        # Normalize line endings, then split on blank lines
        text = re.sub(r"\r\n", "\n", text)
        paragraphs = re.split(r"\n{2,}", text)
        return [p.replace("\n", " ").strip() for p in paragraphs]

    def _looks_like_heading(self, text: str) -> bool:
        """Heuristic: short lines in title case or all caps are headings."""
        stripped = text.strip()
        if len(stripped) > 120:
            return False
        if stripped.isupper() and len(stripped) > 3:
            return True
        words = stripped.split()
        if len(words) <= 10 and sum(1 for w in words if w[0].isupper()) > len(words) * 0.7:
            return True
        return bool(re.match(r"^\d+[\.\)]\s+\w", stripped))

    def _table_to_markdown(self, table: list[list]) -> str:
        """Convert a 2D table to Markdown pipe format."""
        if not table:
            return ""
        # Filter None cells
        table = [[str(cell or "").strip() for cell in row] for row in table]
        table = [row for row in table if any(cell for cell in row)]
        if not table:
            return ""

        header = "| " + " | ".join(table[0]) + " |"
        separator = "| " + " | ".join(["---"] * len(table[0])) + " |"
        rows = ["| " + " | ".join(row) + " |" for row in table[1:]]

        return "\n".join([header, separator] + rows)